"""
Word document extractor for .docx and .doc files
"""
import time
import io
from typing import List
from .base_extractor import BaseExtractor, ExtractionResult


class WordExtractor(BaseExtractor):
    """Word document text extraction"""
    
    @property
    def supported_mime_types(self) -> List[str]:
        return [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',  # .docx
            'application/msword',  # .doc
            'application/vnd.ms-word'
        ]
    
    @property
    def supported_extensions(self) -> List[str]:
        return ['docx', 'doc']
    
    def extract_text(self, file_data: bytes, filename: str = "") -> ExtractionResult:
        """
        Extract text from Word documents
        
        Args:
            file_data: Word file bytes
            filename: Original filename
            
        Returns:
            ExtractionResult with extracted text and metadata
        """
        start_time = time.time()
        
        # Validate file size
        self.validate_file_size(file_data, max_size_mb=100)
        
        # Determine Word format
        is_docx = filename.lower().endswith('.docx') or b'PK' in file_data[:10]
        
        try:
            if is_docx:
                return self._extract_docx(file_data, filename, start_time)
            else:
                return self._extract_doc(file_data, filename, start_time)
                
        except Exception as e:
            raise Exception(f"Word document processing failed: {str(e)}")
    
    def _extract_docx(self, file_data: bytes, filename: str, start_time: float) -> ExtractionResult:
        """Extract text from .docx files using python-docx"""
        try:
            from docx import Document
            from docx.oxml.exceptions import InvalidXmlError
            
            # Load document from bytes
            doc = Document(io.BytesIO(file_data))
            
            all_text = ""
            paragraph_count = 0
            table_count = 0
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    all_text += para_text + "\n"
                    paragraph_count += 1
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_text = f"\n[Table {table_idx + 1}]\n"
                table_has_data = False
                
                for row in table.rows:
                    row_data = []
                    row_has_data = False
                    
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                            row_has_data = True
                        else:
                            row_data.append("")
                    
                    if row_has_data:
                        table_text += " | ".join(row_data) + "\n"
                        table_has_data = True
                
                if table_has_data:
                    all_text += table_text
                    table_count += 1
            
            processing_time = time.time() - start_time
            
            # Extract core properties if available
            metadata = {
                'file_type': 'word_docx',
                'filename': filename,
                'paragraphs': paragraph_count,
                'tables': table_count,
                'processing_time_seconds': processing_time,
                'extractor': 'python-docx'
            }
            
            try:
                core_props = doc.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else ''
                })
            except Exception:
                pass  # Ignore metadata extraction errors
            
            print("Word (.docx) processing completed:")
            print(f"  - Paragraphs: {paragraph_count}")
            print(f"  - Tables: {table_count}")
            print(f"  - Text length: {len(all_text)} characters")
            print(f"  - Processing time: {processing_time:.2f}s")
            
            return ExtractionResult(
                text=all_text,
                pages=1,  # Word docs don't have explicit pages in extraction
                metadata=metadata,
                processing_time=processing_time
            )
            
        except ImportError:
            raise Exception("python-docx not installed - cannot process .docx files")
        except InvalidXmlError:
            raise Exception("Invalid or corrupted Word document")
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def _extract_doc(self, file_data: bytes, filename: str, start_time: float) -> ExtractionResult:
        """Extract text from .doc files using docx2txt as fallback"""
        try:
            import docx2txt
            
            # docx2txt can handle some .doc files, but it's primarily for .docx
            # For .doc files, we'll try a simple approach first
            
            # Save to temporary file and extract
            import tempfile
            import os
            
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                temp_file.write(file_data)
                temp_file.flush()
                
                try:
                    # Try extracting with docx2txt
                    extracted_text = docx2txt.process(temp_file.name)
                    
                    processing_time = time.time() - start_time
                    
                    if extracted_text and extracted_text.strip():
                        all_text = extracted_text.strip()
                        
                        metadata = {
                            'file_type': 'word_doc',
                            'filename': filename,
                            'processing_time_seconds': processing_time,
                            'extractor': 'docx2txt',
                            'note': 'Legacy .doc format - limited extraction capabilities'
                        }
                        
                        print("Word (.doc) processing completed:")
                        print(f"  - Text length: {len(all_text)} characters") 
                        print(f"  - Processing time: {processing_time:.2f}s")
                        print("  - Note: Legacy format with limited extraction")
                        
                        return ExtractionResult(
                            text=all_text,
                            pages=1,
                            metadata=metadata,
                            processing_time=processing_time
                        )
                    else:
                        raise Exception("No text extracted from .doc file")
                        
                finally:
                    # Clean up temporary file
                    try:
                        os.unlink(temp_file.name)
                    except Exception:
                        pass
            
        except ImportError:
            raise Exception("docx2txt not installed - cannot process .doc files")
        except Exception as e:
            # Fallback: return minimal result indicating unsupported format
            processing_time = time.time() - start_time
            
            metadata = {
                'file_type': 'word_doc',
                'filename': filename,
                'processing_time_seconds': processing_time,
                'extractor': 'none',
                'error': str(e),
                'note': 'Legacy .doc format not fully supported - consider converting to .docx'
            }
            
            print(f"Word (.doc) processing failed: {str(e)}")
            print("  - Recommendation: Convert .doc files to .docx format for better support")
            
            return ExtractionResult(
                text=f"[Unable to extract text from legacy .doc file: {filename}]\n[Error: {str(e)}]\n[Recommendation: Please convert to .docx format]",
                pages=1,
                metadata=metadata,
                processing_time=processing_time
            )